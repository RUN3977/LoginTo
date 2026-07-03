from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "LoginTo_MVP与实施规划.docx"

COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "muted": "667085",
    "light_gray": "F2F4F7",
    "callout": "F4F6F9",
    "border": "D0D5DD",
}


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_run(run, bold=False, size=10.5, color="000000"):
    set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, after=6, size=10.8):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    style_run(run, size=size)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        style_run(run)


def add_numbers(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        style_run(run)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["border"]):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def style_cell_text(cell, bold=False, color="000000", size=10):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            style_run(run, bold=bold, color=color, size=size)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    mark_row_as_header(table.rows[0])

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, COLORS["light_gray"])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        style_cell_text(cell, bold=True, color=COLORS["ink"])
        if widths:
            set_cell_width(cell, widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            style_cell_text(cells[idx])
            if widths:
                set_cell_width(cells[idx], widths[idx])

    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.12
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), COLORS["callout"])
    p_pr.append(shd)

    title_run = paragraph.add_run(title + " ")
    style_run(title_run, bold=True, color=COLORS["dark_blue"])
    body_run = paragraph.add_run(body)
    style_run(body_run, color=COLORS["ink"])


def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)

    for line_idx, line in enumerate(text.splitlines()):
        if line_idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string("24292F")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(COLORS["ink"])

    styles["Subtitle"].font.name = "Calibri"
    styles["Subtitle"].font.size = Pt(11)
    styles["Subtitle"].font.color.rgb = RGBColor.from_string(COLORS["muted"])

    heading_tokens = {
        "Heading 1": (16, COLORS["blue"], 16, 8),
        "Heading 2": (13, COLORS["blue"], 12, 6),
        "Heading 3": (12, COLORS["dark_blue"], 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("LoginTo MVP 规划 | 终端设备版 v0.2 | 2026-06-05")
    style_run(run, size=9, color=COLORS["muted"])


def add_title_page(doc):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("LoginTo 多终端近场同步账号记录 App")
    style_run(run, bold=True, size=24, color=COLORS["ink"])

    subtitle = doc.add_paragraph(style="Subtitle")
    run = subtitle.add_run("MVP 与实施规划 · 终端设备版 · 版本 v0.2 · 2026-06-05")
    style_run(run, size=11, color=COLORS["muted"])

    add_callout(
        doc,
        "核心判断：",
        "LoginTo 不再规划网页/PWA 端。MVP 聚焦移动终端、桌面终端和平板/备用设备场景，"
        "通过本地加密、离线优先、面对面近场同步来保护账号、银行卡、会员、证件和密钥信息。"
    )

    add_para(
        doc,
        "本文档用于锁定开始写代码前的产品范围、技术路线、阶段计划和验收标准。"
        "后续实现应优先遵守本文档，只有在安全性、平台能力或终端体验出现明确冲突时才调整。"
    )


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("1. 产品定位", level=1)
    add_para(
        doc,
        "LoginTo 是一个用于记录社交账号、网站账号、银行卡、会员信息、证件、密钥和自定义私密资料的多终端应用。"
        "它的差异点不是“随处云同步”，而是“用户亲自控制数据在哪里、什么时候、怎样同步”。"
    )
    add_bullets(doc, [
        "本地优先：所有记录、附件、提醒和同步日志默认保存在本机加密保险库中。",
        "无公共云登录：MVP 不引入中心账号系统，也不将用户保险库上传到公共服务器。",
        "终端近场同步：手机、平板、电脑、备用机之间通过扫码配对、局域网、热点或加密同步包传输。",
        "智能整理：拍照或导入图片后，本地 OCR 自动提取字段并生成待确认记录。",
        "提醒可靠：银行卡、会员、证件、合同、订阅等记录可以设置到期弹窗和本地系统通知。",
    ])

    doc.add_heading("2. 产品原则", level=1)
    add_table(
        doc,
        ["原则", "含义", "实现约束"],
        [
            ["安全可控优先", "敏感内容不依赖第三方云端可用性。", "不做云端明文、不做云端 OCR、不做默认联网分析。"],
            ["自动化但不替用户决定", "OCR 和分类可以自动建议，但入库前必须确认。", "智能录入生成草稿，用户确认后才写入正式记录。"],
            ["同步过程可解释", "用户应知道正在和哪台终端同步、同步了什么。", "配对验证码、设备名称、同步摘要和冲突列表必须可见。"],
            ["记录模型可扩展", "先覆盖常见类型，之后可以添加模板。", "以通用字段模型承载账号、银行卡、会员、证件等差异。"],
            ["离线可用", "没有互联网时仍能查看、编辑、提醒和备份。", "所有核心能力均在本地运行，联网只用于用户主动近场传输。"],
        ],
        widths=[2100, 3500, 3760],
    )

    doc.add_heading("3. 终端设备形态", level=1)
    add_table(
        doc,
        ["终端", "MVP 角色", "关键能力", "同步方式"],
        [
            ["移动终端", "快速查询、拍照录入、提醒触达。", "主密码/生物识别解锁、相机 OCR、本地通知、扫码配对。", "二维码配对后通过局域网/热点同步。"],
            ["平板终端", "大屏整理、附件校对、家庭备用设备。", "继承移动端能力，强化列表、详情和批量确认体验。", "与手机同属移动终端同步通道。"],
            ["桌面终端", "批量整理、稳定同步中枢、本地管理。", "记录管理、批量编辑、附件查看、托盘、局域网同步服务、加密备份。", "同 Wi-Fi 发现、手动 IP、手机热点。"],
            ["新/备用设备", "换机、恢复、离线灾备。", "导入加密备份包、面对面配对、恢复提醒与附件。", "旧设备扫码授权或手动导入加密备份。"],
        ],
        widths=[1450, 2450, 3300, 2160],
    )
    add_callout(
        doc,
        "网页端决策：",
        "MVP 不做网页/PWA，不提供浏览器里的账号管理页面。桌面端可以使用 React 作为界面技术，"
        "但它以 Tauri 本地应用发布，拥有本地文件、通知、托盘、设备密钥和局域网服务能力。"
    )

    doc.add_heading("4. MVP 范围", level=1)
    doc.add_heading("4.1 MVP 必须包含", level=2)
    add_bullets(doc, [
        "创建本地保险库：主密码初始化、解锁、自动锁定、错误次数保护。",
        "记录管理：账号、银行卡、会员、证件、密钥、自定义记录的新增、编辑、删除、收藏、归档。",
        "敏感字段：密码、卡号、证件号、密钥默认隐藏，复制后自动清空剪贴板。",
        "分类与标签：内置分类、用户自定义标签、全局搜索和最近使用。",
        "提醒系统：到期日、提前提醒、弹窗内容、稍后提醒、已处理状态。",
        "图片附件：图片加密保存，可作为记录附件或备注依据。",
        "拍照自动整理：本地 OCR、类型判断、字段提取、确认后入库。",
        "面对面同步：扫码配对、加密通道、变更同步、冲突保留和手动合并。",
        "加密备份：导出/导入加密保险库包，用于换机、恢复和网络受限环境。",
    ])

    doc.add_heading("4.2 MVP 暂不包含", level=2)
    add_bullets(doc, [
        "网页/PWA、网页管理端、浏览器内保险库。",
        "公共云账户、云端自动同步、云端备份。",
        "浏览器自动填充插件和网页表单自动登录。",
        "团队空间、多人权限、远程共享。",
        "云 OCR、云 AI 分类、第三方分析埋点。",
        "银行卡支付能力、余额查询、交易记录同步。",
        "自动抓取短信验证码或邮件内容。",
    ])

    doc.add_heading("5. 记录类型与字段模型", level=1)
    add_table(
        doc,
        ["记录类型", "核心字段", "提醒场景", "特别规则"],
        [
            ["网站/社交账号", "用户名、密码、邮箱、手机号、网址、2FA 备用码、备注。", "密码轮换、会员到期、账号年审。", "密码和 2FA 默认高敏隐藏。"],
            ["银行卡", "银行、卡号、持卡人、有效期、账单日、还款日、预留手机号。", "有效期、年费日、账单日、还款日。", "不默认记录 CVV；如用户强制保存，应二次解锁。"],
            ["会员信息", "会员名、会员号、等级、权益、到期日、客服电话。", "到期续费、权益变更、积分清零。", "到期提醒默认提前 7 天。"],
            ["证件信息", "证件类型、号码、签发地、有效期、附件照片。", "证件过期、签证过期、年检。", "证件照片作为高敏附件保存。"],
            ["密钥/API", "名称、平台、Key/Secret、用途、权限范围、创建时间。", "密钥轮换、权限复查。", "Secret 字段永不出现在搜索明文缓存。"],
            ["自定义记录", "标题、动态字段、标签、附件、备注。", "用户自定义日期字段提醒。", "字段可标记为普通/敏感/高敏。"],
        ],
        widths=[1500, 3500, 2300, 2060],
    )

    doc.add_heading("6. 提醒系统设计", level=1)
    add_para(
        doc,
        "提醒是本地能力，不依赖服务器。每条记录可以拥有多个提醒规则；提醒触发后显示系统通知和应用内弹窗，"
        "用户可以选择已处理、稍后提醒或打开对应记录。"
    )
    add_table(
        doc,
        ["能力", "MVP 设定", "验收标准"],
        [
            ["到期提醒", "支持当天、提前 1/3/7/30 天和自定义天数。", "关闭 App 后，移动终端和桌面终端仍能按平台能力触发通知。"],
            ["弹窗内容", "用户可为每条提醒写自定义提示。", "弹窗显示标题、类型、到期日、自定义内容和操作按钮。"],
            ["重复提醒", "支持一次性、每天、每周、每月、每年。", "重复规则可编辑，已处理后按下一次规则计算。"],
            ["稍后提醒", "支持 10 分钟、1 小时、明天、自定义时间。", "稍后提醒不会修改原始到期日期。"],
            ["提醒审计", "记录提醒创建、触发、稍后、完成。", "同步后其他终端能看到提醒状态变化。"],
        ],
        widths=[1600, 4200, 3560],
    )

    doc.add_heading("7. 拍照自动整理设计", level=1)
    add_numbers(doc, [
        "用户拍照或导入图片，图片先进入本地临时处理区。",
        "本地 OCR 提取文字，不上传到公共云服务。",
        "分类器根据关键词、数字结构、日期格式和版面特征判断记录类型。",
        "字段提取器生成草稿，例如会员号、卡号后四位、到期日、客服电话、平台名。",
        "用户在确认页检查字段、修改分类、选择是否设置提醒。",
        "确认后创建记录，并将原图作为加密附件保存到记录备注区。",
    ])
    add_callout(
        doc,
        "智能录入底线：",
        "OCR 只能生成草稿，不能直接写入正式保险库；无法识别的字段保留在 OCR 文本备注中，避免误删信息。"
    )

    doc.add_heading("8. 本地数据与加密设计", level=1)
    add_table(
        doc,
        ["组件", "职责", "MVP 选型建议"],
        [
            ["保险库文件", "承载记录、字段、索引、提醒和同步日志。", "SQLite 或嵌入式数据库；记录级应用层加密。"],
            ["附件仓库", "保存照片、截图、备份码图片等大文件。", "独立加密 blob 文件，元数据写入数据库。"],
            ["主密码", "解锁保险库，不上传、不保存明文。", "Argon2id 派生密钥；参数随设备性能校准。"],
            ["保险库密钥", "真正加密数据的随机密钥。", "由主密码派生密钥包裹保存，后续支持设备密钥包裹。"],
            ["加密算法", "保护记录内容和附件。", "优先 libsodium/XChaCha20-Poly1305；按桌面/移动平台封装统一接口。"],
            ["搜索索引", "加速本地搜索。", "解锁后在内存构建；敏感字段默认不进入可检索明文索引。"],
            ["设备密钥", "识别和配对终端。", "设备生成公私钥对，私钥放入系统安全存储。"],
        ],
        widths=[1500, 3800, 4060],
    )

    doc.add_heading("9. 面对面同步设计", level=1)
    add_numbers(doc, [
        "终端 A 显示配对二维码，包含设备 ID、公钥、一次性会话码和局域网入口。",
        "终端 B 扫码后展示双方设备名称与 6 位验证码，用户面对面核对。",
        "双方完成密钥协商，建立加密通道。",
        "同步模块交换变更摘要，只传输对方缺少的变更包和附件块。",
        "合并时按记录版本、字段修改时间和设备变更日志处理；冲突不覆盖，生成待合并版本。",
        "同步完成后展示新增、更新、冲突、附件数量和失败项。",
    ])
    add_table(
        doc,
        ["传输方式", "MVP 优先级", "适用场景", "注意事项"],
        [
            ["同 Wi-Fi 局域网", "高", "手机、平板和电脑在同一网络。", "需要 mDNS/手动 IP 兜底。"],
            ["手机热点", "高", "没有可信 Wi-Fi 时面对面同步。", "桌面终端连接手机热点后同步。"],
            ["加密同步包", "高", "换机、备用机、网络受限环境。", "需要用户手动导出和导入，作为近场同步兜底。"],
            ["蓝牙/BLE", "中", "小数据量或配对辅助。", "跨平台稳定性和速度需验证。"],
            ["私有中继/NAS", "低", "未来高级用户自建同步。", "只作为可选能力，不进入 MVP。"],
        ],
        widths=[1700, 1200, 3200, 3260],
    )

    doc.add_heading("10. 数据模型初稿", level=1)
    add_table(
        doc,
        ["实体", "关键字段", "说明"],
        [
            ["vaults", "id, name, schema_version, created_at", "本地保险库元信息。"],
            ["records", "id, type, title, category_id, favorite, archived, version", "记录主体，不直接保存明文字段。"],
            ["record_fields", "record_id, key, label, value_cipher, sensitivity", "动态字段，支持普通/敏感/高敏。"],
            ["categories", "id, name, icon, sort_order", "内置分类和用户自定义分类。"],
            ["tags", "id, name, color", "跨分类组织。"],
            ["attachments", "id, record_id, blob_path, mime, digest, encrypted_size", "图片和文件附件元数据。"],
            ["reminders", "id, record_id, due_at, rule, message, status", "到期、重复、稍后提醒。"],
            ["devices", "id, name, public_key, trusted_at, last_seen_at", "已配对终端设备。"],
            ["sync_changes", "id, entity, entity_id, op, device_id, lamport, created_at", "同步变更日志。"],
            ["conflicts", "id, record_id, local_version, remote_version, status", "冲突版本和合并状态。"],
        ],
        widths=[1700, 3500, 4160],
    )

    doc.add_heading("11. 技术架构建议", level=1)
    add_table(
        doc,
        ["层", "建议", "原因"],
        [
            ["桌面终端", "Tauri + React + TypeScript", "体积小、系统能力强，适合本地文件、托盘、通知和局域网服务。"],
            ["移动/平板终端", "React Native / Expo Dev Client", "共享 TypeScript 业务逻辑，同时使用相机、通知和生物识别能力。"],
            ["核心包", "packages/vault-core, sync-core, crypto-core, ocr-core", "把保险库、同步、加密、OCR 流程从 UI 中分离。"],
            ["数据库", "SQLite + 应用层字段加密", "跨桌面/移动可用，便于变更日志和本地查询。"],
            ["测试", "Vitest + Tauri 集成测试 + 移动端集成测试", "核心逻辑可单元测试，关键 UI 和同步流程需要端到端验证。"],
        ],
        widths=[1600, 3300, 4460],
    )
    add_code_block(doc, """apps/
  desktop/        Tauri 桌面终端应用
  mobile/         React Native 移动/平板终端应用
packages/
  vault-core/     记录模型、保险库读写、迁移
  crypto-core/    密钥派生、加密封装、敏感字段策略
  sync-core/      设备配对、变更日志、冲突合并
  ocr-core/       OCR 草稿、字段提取、分类建议
  ui/             跨终端组件和设计系统
docs/
  LoginTo_MVP与实施规划.docx""")

    doc.add_heading("12. MVP 开发进程设定", level=1)
    add_table(
        doc,
        ["阶段", "目标", "主要交付", "完成标准"],
        [
            ["M0 开工准备", "锁定范围、技术选择和安全边界。", "项目结构、威胁模型、字段模板、UI 信息架构。", "没有未决的 MVP 阻塞决策。"],
            ["M1 本地保险库", "完成本地加密数据底座。", "主密码、加密记录、附件仓库、迁移机制。", "无网络时可创建、解锁、保存、恢复记录。"],
            ["M2 记录管理", "完成核心使用闭环。", "分类、标签、搜索、收藏、详情页、编辑页。", "账号/银行卡/会员/证件/密钥均可录入和检索。"],
            ["M3 提醒系统", "完成到期提醒闭环。", "提醒规则、本地通知、弹窗、稍后提醒、已处理状态。", "会员和银行卡提醒能在目标终端按时触发。"],
            ["M4 智能录入", "完成拍照到记录草稿。", "相机导入、本地 OCR、类型识别、字段草稿、附件保存。", "用户确认后生成记录，原图作为附件。"],
            ["M5 近场同步", "完成两台终端面对面同步。", "二维码配对、加密通道、变更同步、附件同步、冲突列表。", "手机与电脑能在无公共互联网账号下互相同步。"],
            ["M6 多终端发布候选", "完成 MVP 发布候选。", "桌面终端、移动终端、平板适配、加密备份、验收测试。", "核心数据模型在多终端一致，数据可近场互通。"],
        ],
        widths=[1600, 2500, 3300, 1960],
    )

    doc.add_heading("13. MVP 之后的路线", level=1)
    add_table(
        doc,
        ["版本", "主题", "新增能力"],
        [
            ["v1.1", "同步体验增强", "多终端同步拓扑、同步历史、设备信任等级、蓝牙辅助配对、冲突合并 UI 优化。"],
            ["v1.2", "智能整理增强", "更多卡片/证件模板、批量图片识别、字段置信度、自动标签、附件文字全文检索。"],
            ["v1.3", "终端效率增强", "桌面快捷键、托盘快速搜索、移动端小组件、系统分享入口、批量确认工作台。"],
            ["v1.4", "家庭/小组本地共享", "面对面共享单条记录或共享保险库，按设备授权，不走云端账号体系。"],
            ["v1.5", "高级安全", "硬件安全密钥、紧急恢复包、主密码轮换、导出水印、敏感操作审计。"],
            ["v2.0", "成熟离线生态", "跨平台自动发现更稳、局域网多设备增量同步、可选私有 NAS/自建同步中继。"],
        ],
        widths=[1100, 2200, 6060],
    )

    doc.add_heading("14. MVP 验收标准", level=1)
    add_bullets(doc, [
        "不登录公共云账号，也能完成创建保险库、录入记录、设置提醒、拍照整理和导出备份。",
        "移动终端和桌面终端可以面对面配对并完成加密同步，同步过程有设备确认和结果摘要。",
        "平板终端能复用移动端能力，完成查看、编辑、附件确认和提醒查看。",
        "会员信息和银行卡至少支持一个到期/还款/续费提醒，并能显示自定义弹窗内容。",
        "拍照后能够自动生成记录草稿，并把原图作为加密附件保存。",
        "敏感字段默认隐藏，复制后剪贴板自动清除，高敏字段可要求二次验证。",
        "断网状态下核心功能可用；核心应用不得向公共服务上传保险库、附件或 OCR 内容。",
        "冲突不会静默覆盖，必须保留冲突版本并提示用户合并。",
        "加密备份可以在新设备恢复，恢复后记录、附件、提醒和标签完整。",
    ])

    doc.add_heading("15. 关键风险与处理", level=1)
    add_table(
        doc,
        ["风险", "影响", "处理策略"],
        [
            ["局域网发现不稳定", "不同路由器或系统权限可能导致设备无法自动发现。", "提供扫码入口、手动 IP、手机热点和加密同步包兜底。"],
            ["终端平台能力差异", "通知、生物识别、文件权限、后台运行在各平台表现不同。", "用平台适配层封装能力，设置页显示权限状态和测试入口。"],
            ["OCR 准确率不稳定", "银行卡、证件和会员卡字段可能识别错误。", "只生成草稿，显示置信度，用户确认后入库。"],
            ["主密码丢失", "本地加密保险库无法恢复。", "MVP 提供恢复说明和加密备份；后续支持紧急恢复包。"],
            ["多终端冲突", "同时编辑同一记录可能覆盖信息。", "使用变更日志和冲突表，默认保留双方版本。"],
            ["附件体积过大", "同步慢、备份包膨胀。", "附件分块、压缩缩略图、按需同步原图。"],
        ],
        widths=[2500, 3100, 3760],
    )

    doc.add_heading("16. 开始写代码前清单", level=1)
    add_numbers(doc, [
        "确认终端定义：移动终端、平板终端、桌面终端、新/备用设备；不做网页/PWA。",
        "确认技术栈：Tauri + React、React Native、TypeScript monorepo。",
        "确认加密库：优先 libsodium，并按桌面/移动平台封装统一 crypto-core。",
        "确认 OCR 路径：移动端优先系统/ML Kit 本地 OCR，桌面端先接 Tesseract 或系统 OCR。",
        "画出第一版信息架构：保险库解锁、首页、列表、详情、编辑、提醒、同步、设置。",
        "定义记录模板 JSON Schema 和数据库迁移 v1。",
        "实现威胁模型文档，覆盖丢设备、导出泄露、剪贴板泄露、同步中间人、恶意附件。",
        "搭建测试策略：crypto-core 和 sync-core 必须先有单元测试，再做 UI。",
        "建立样例数据：社交账号、网站账号、银行卡、会员卡、证件、API Key、带附件记录。",
        "M0 全部通过后再正式进入代码实现。",
    ])

    doc.add_heading("17. 默认决策记录", level=1)
    add_table(
        doc,
        ["决策", "当前结论", "后续可调整条件"],
        [
            ["同步路线", "不做公共云登录；默认面对面近场同步。", "只有用户明确要求私有云/NAS 时才加入可选中继。"],
            ["网页端", "不做网页/PWA，也不把浏览器作为主要使用入口。", "除非产品方向被重新定义，否则不纳入 MVP 或近期路线。"],
            ["桌面界面", "可使用 React 技术，但必须作为 Tauri 本地终端应用发布。", "不能因为用了前端技术而削弱本地文件、通知、设备密钥和局域网能力。"],
            ["智能整理", "OCR 全本地，结果入库前必须确认。", "若未来允许云模型，也必须是显式开关。"],
            ["银行卡 CVV", "不默认保存。", "如保存，需高敏字段、二次解锁和强提醒。"],
            ["发布顺序", "先桌面 + 移动打通；平板随移动端适配。", "如果实际测试表明平板录入体验更关键，可提高平板优化优先级。"],
        ],
        widths=[1800, 4300, 3260],
    )

    doc.core_properties.title = "LoginTo MVP 与实施规划"
    doc.core_properties.subject = "多终端近场同步账号记录 App 开工前规划"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "本地优先、端到端加密、面对面同步的账号记录 App 规划文档。"

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_document())
